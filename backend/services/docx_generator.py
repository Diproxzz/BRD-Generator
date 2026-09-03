import os
from typing import Dict, Any, List
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PRIMARY_COLOR = RGBColor(26, 54, 93)     # #1A365D Deep Navy
SECONDARY_COLOR = RGBColor(49, 130, 206) # #3182CE Corporate Blue
TEXT_COLOR = RGBColor(45, 55, 72)        # #2D3748 Slate Charcoal
LIGHT_BG_HEX = "EDF2F7"                  # Light grey/blue for table headers
BORDER_HEX = "CBD5E0"

def set_cell_shading(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table, color_hex=BORDER_HEX):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def format_cell_text(cell, text: str, bold: bool = False, color: RGBColor = TEXT_COLOR, font_size: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(str(text if text is not None else ""))
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color

def add_styled_table(doc: Document, headers: List[str], rows: List[List[Any]], col_widths: List[float] = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header row
    hdr_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_shading(cell, LIGHT_BG_HEX)
        format_cell_text(cell, header, bold=True, color=PRIMARY_COLOR, font_size=10)
        
    # Content rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[1 + r_idx]
        # Alternating background
        row_bg = "FFFFFF" if r_idx % 2 == 0 else "F7FAFC"
        for c_idx, cell_value in enumerate(row_data):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                if row_bg != "FFFFFF":
                    set_cell_shading(cell, row_bg)
                format_cell_text(cell, str(cell_value), bold=False, color=TEXT_COLOR, font_size=9.5)
                
    if col_widths and len(col_widths) == len(headers):
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_heading_styled(doc: Document, text: str, level: int):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
    else:
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(2)
    return heading

def build_docx_brd(data: Dict[str, Any], output_path: str) -> str:
    doc = Document()
    
    # Page Margins (1 inch everywhere)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    project_name = data.get("project_name", "Enterprise Data Pipeline")
    version = data.get("version", "1.0")
    date_str = data.get("date", "2026-09-03")
    author = data.get("author", "Business Analyst")
    
    # COVER / TITLE BLOCK
    doc.add_paragraph().paragraph_format.space_before = Pt(36)
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(project_name)
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    
    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Business Requirements Document (BRD)")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = SECONDARY_COLOR
    subtitle_p.paragraph_format.space_after = Pt(24)
    
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Document Version: {version}\nDate: {date_str}\nAuthor: {author}\nStatus: Approved for Review").font.color.rgb = TEXT_COLOR
    meta_p.paragraph_format.space_after = Pt(36)
    
    doc.add_page_break()
    
    # 1. Revision History
    add_heading_styled(doc, "1. Revision History", 1)
    rev_headers = ["Version Number", "Date", "Author", "Description"]
    rev_rows = data.get("revision_history", [
        ["1.0", date_str, author, "Initial comprehensive draft generated from source documentation"]
    ])
    add_styled_table(doc, rev_headers, rev_rows, [1.2, 1.2, 1.6, 2.5])
    
    # 2. Project Overview
    add_heading_styled(doc, "2. Project Overview", 1)
    
    # 2.1 Project Sponsor(s)
    add_heading_styled(doc, "2.1 Project Sponsor(s)", 2)
    sponsors = data.get("sponsors", [["[NEEDS INPUT: Project Sponsor]", "[NEEDS INPUT: Job Title]"]])
    add_styled_table(doc, ["Name", "Job Title"], sponsors, [3.2, 3.3])
    
    # 2.2 Project Contributors
    add_heading_styled(doc, "2.2 Project Contributors (A–Z)", 2)
    contributors = data.get("contributors", [
        [author, "Lead Business Analyst", "Requirements Gathering & Documentation"],
        ["[NEEDS INPUT: Technical Lead]", "Solutions Architect", "Technical Architecture"],
        ["[NEEDS INPUT: Product Owner]", "Product Manager", "Scope & Acceptance Criteria"]
    ])
    add_styled_table(doc, ["Name", "Job Title", "Role"], contributors, [2.2, 2.2, 2.1])
    
    # 2.3 In Scope
    add_heading_styled(doc, "2.3 In Scope (Deliverables)", 2)
    in_scope = data.get("in_scope", [["Automated pipeline ingestion"], ["Real-time data validation"]])
    add_styled_table(doc, ["Title"], in_scope, [6.5])
    
    # 2.4 Out of Scope
    add_heading_styled(doc, "2.4 Out of Scope", 2)
    out_of_scope = data.get("out_of_scope", [
        ["Legacy batch export module", "Deprecated system to be retired in Q4"],
        ["Third-party customer billing UI", "Managed by external payment vendor"]
    ])
    add_styled_table(doc, ["Title", "Reason for Exclusion"], out_of_scope, [3.2, 3.3])
    
    # 3. Acronyms
    add_heading_styled(doc, "3. Common Project Acronyms, Names, and Descriptions", 1)
    acronyms = data.get("acronyms", [
        ["BRD", "Business Requirements Document"],
        ["API", "Application Programming Interface"],
        ["ETL", "Extract, Transform, Load"],
        ["SLA", "Service Level Agreement"]
    ])
    add_styled_table(doc, ["Name", "Description"], acronyms, [2.0, 4.5])
    
    # 4. Existing Processes
    add_heading_styled(doc, "4. Existing Processes", 1)
    proc = data.get("existing_processes", {})
    add_heading_styled(doc, "4.1 Summary Process Narrative", 2)
    doc.add_paragraph(proc.get("summary", "The current business process relies on manual spreadsheet exports and scheduled email attachments.")).paragraph_format.space_after = Pt(4)
    
    add_heading_styled(doc, "4.2 Timing", 2)
    doc.add_paragraph(proc.get("timing", "Executes daily at 06:00 UTC with manual review cycles taking up to 4 hours.")).paragraph_format.space_after = Pt(4)
    
    add_heading_styled(doc, "4.3 Volume", 2)
    doc.add_paragraph(proc.get("volume", "Approximately 50,000 records processed per business day.")).paragraph_format.space_after = Pt(4)
    
    add_heading_styled(doc, "4.4 Screenshots / Flow Diagrams", 2)
    doc.add_paragraph(proc.get("screenshots", "Architectural and UI references uploaded during initial discovery.")).paragraph_format.space_after = Pt(4)
    
    add_heading_styled(doc, "4.5 Problems", 2)
    doc.add_paragraph(proc.get("problems", "High manual latency, lack of validation checks, error-prone data transformations, and lack of real-time auditability.")).paragraph_format.space_after = Pt(4)
    
    # 5. Project Requirements
    add_heading_styled(doc, "5. Project Requirements", 1)
    deliverables = data.get("deliverables", [])
    if not deliverables:
        deliverables = [{
            "title": "Data Pipeline Core Workflow",
            "process_overview": {
                "summary": "Automate data intake, schema enforcement, and publish sanitized records.",
                "flow_diagram": "Source Upload -> Virus & Format Scan -> Schema Validation -> Transform -> Storage -> Notification",
                "trigger": "Incoming payload received via secure webhook or sFTP batch drop.",
                "timing": "Near real-time execution with latency under 1500ms per record batch.",
                "volume": "Designed to scale up to 1,000,000 records/day.",
                "outcomes": "Records validated, transformed, persisted to data warehouse, and status webhook broadcasted."
            },
            "functional_requirements": [
                {"id": "PREQ-001", "text": "The system shall ingest structured records from configured source endpoints.", "level": "PREQ"},
                {"id": "CREQ-001.1", "text": "The ingestion gateway must validate payload size against max allowable limit of 25MB.", "level": "CREQ"},
                {"id": "GCREQ-001.1.1", "text": "Paylod rejects over 25MB must trigger an HTTP 413 Payload Too Large response with structured error code.", "level": "GCREQ"},
                {"id": "CREQ-001.2", "text": "The gateway must authenticate all API requests via OAuth2 Bearer token.", "level": "CREQ"},
                {"id": "PREQ-002", "text": "The system shall parse and normalize disparate date and currency attributes.", "level": "PREQ"},
                {"id": "CREQ-002.1", "text": "All timestamps must be standardized to ISO 8601 UTC format.", "level": "CREQ"}
            ],
            "non_functional_requirements": {
                "Availability": "PREQ-NFR-001: The system shall maintain 99.95% operational uptime during standard business calendar.",
                "Compatibility": "PREQ-NFR-002: Must support REST JSON payloads, Parquet, and CSV standard RFC 4180 formats.",
                "Extensibility": "PREQ-NFR-003: Modular parser plugins allowing new source format adapters without service restart.",
                "Maintainability": "PREQ-NFR-004: All code modules must maintain minimum 80% automated unit test coverage.",
                "Scalability": "PREQ-NFR-005: Horizontal auto-scaling of consumer workers based on queue depth metrics.",
                "Security": "PREQ-NFR-006: Encryption at rest using AES-256 and in transit via TLS 1.3.",
                "Usability": "PREQ-NFR-007: Administrative dashboard must provide self-service error search and replay in under 3 clicks.",
                "Performance": "PREQ-NFR-008: 95th percentile pipeline processing latency must not exceed 2.0 seconds per batch."
            },
            "data_requirements": [
                ["record_id", "Unique GUID assigned to incoming event", "No", "Yes", "UUIDv4"],
                ["source_system", "Originating client or application identifier", "No", "Yes", "Alpha-numeric (max 32)"],
                ["payload_content", "Sanitized record JSON payload", "Yes", "Yes", "Valid JSON structure"],
                ["created_timestamp", "System receipt timestamp in UTC", "No", "Yes", "ISO 8601"]
            ],
            "risks_and_assumptions": [
                ["Assumption", "Client applications provide standard OAuth2 bearer tokens."],
                ["Risk", "Source systems may experience sporadic network latency during peak hours."],
                ["Dependency", "Dependent on Cloud Storage availability for cold archiving."]
            ]
        }]
        
    for d_idx, d in enumerate(deliverables, 1):
        d_title = d.get("title", f"Deliverable {d_idx}")
        add_heading_styled(doc, f"5.{d_idx} {d_title}", 2)
        
        # 5.x.1 Process Overview
        add_heading_styled(doc, f"5.{d_idx}.1 Process Overview", 3)
        po = d.get("process_overview", {})
        doc.add_paragraph(f"Summary Process Narrative: {po.get('summary', '[NEEDS INPUT: Narrative]')}")
        doc.add_paragraph(f"Flow Diagram: {po.get('flow_diagram', '[NEEDS INPUT: Diagram flow]')}")
        doc.add_paragraph(f"Triggering Event & Pre-Conditions: {po.get('trigger', '[NEEDS INPUT: Trigger]')}")
        doc.add_paragraph(f"Timing: {po.get('timing', '[NEEDS INPUT: Timing]')}")
        doc.add_paragraph(f"Volume: {po.get('volume', '[NEEDS INPUT: Volume]')}")
        doc.add_paragraph(f"Outcome(s) / Post-Conditions: {po.get('outcomes', '[NEEDS INPUT: Outcomes]')}")
        
        # 5.x.2 Functional Requirements (PREQ/CREQ/GCREQ)
        add_heading_styled(doc, f"5.{d_idx}.2 Functional Requirements", 3)
        frs = d.get("functional_requirements", [])
        for req in frs:
            req_id = req.get("id", "PREQ-000")
            req_text = req.get("text", "")
            level = req.get("level", "PREQ")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            if level == "PREQ":
                p.paragraph_format.left_indent = Inches(0.0)
                p.add_run(f"• {req_id}: ").bold = True
            elif level == "CREQ":
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run(f"  - {req_id}: ").bold = True
            else: # GCREQ
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(f"    * {req_id}: ").bold = True
            p.add_run(req_text)
            
        # 5.x.3 Non-Functional Requirements
        add_heading_styled(doc, f"5.{d_idx}.3 Non-Functional Requirements", 3)
        nfrs = d.get("non_functional_requirements", {})
        nfr_categories = ["Availability", "Compatibility", "Extensibility", "Maintainability", 
                          "Scalability", "Security", "Usability", "Performance"]
        for cat in nfr_categories:
            val = nfrs.get(cat, f"[NEEDS INPUT: {cat} criteria]")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"• {cat}: ")
            run.bold = True
            p.add_run(val)
            
        # 5.x.4 Data Requirements
        add_heading_styled(doc, f"5.{d_idx}.4 Data Requirements", 3)
        data_reqs = d.get("data_requirements", [])
        if data_reqs:
            add_styled_table(
                doc,
                ["Data Field Name", "Description", "Editable", "Mandatory Field", "Predefined Value(s)"],
                data_reqs,
                [1.5, 2.0, 0.8, 1.0, 1.2]
            )
        
        # Known Issues/Assumptions/Risks/Dependencies
        add_heading_styled(doc, f"5.{d_idx}.5 Known Issues, Assumptions, Risks & Dependencies", 3)
        risks = d.get("risks_and_assumptions", [])
        if risks:
            add_styled_table(doc, ["Type", "Description"], risks, [1.8, 4.7])
            
    # 6. Sign off
    add_heading_styled(doc, "6. Sign off", 1)
    sign_headers = ["Project Role", "Signature", "Date"]
    sign_rows = data.get("sign_off", [
        ["Project Sponsor / Functional Lead", "___________________________", "[Pending Sign-off]"],
        ["Technical Lead / Architect", "___________________________", "[Pending Sign-off]"],
        ["Lead Business Analyst", "___________________________", "[Pending Sign-off]"]
    ])
    add_styled_table(doc, sign_headers, sign_rows, [2.5, 2.5, 1.5])
    
    # 7. Appendix
    add_heading_styled(doc, "7. Appendix", 1)
    app = data.get("appendix", {})
    add_heading_styled(doc, "7.1 Mock-ups", 2)
    doc.add_paragraph(app.get("mockups", "Mock-up wireframes and flow visual references documented in discovery assets."))
    
    add_heading_styled(doc, "7.2 Glossary", 2)
    doc.add_paragraph(app.get("glossary", "Definitions of domain terms and pipeline terminology."))
    
    add_heading_styled(doc, "7.3 Business Rules and Procedures", 2)
    doc.add_paragraph(app.get("business_rules", "BR-01: Inactive tenant accounts must be quarantined.\nBR-02: Daily reconciliation batch runs at cutoff window."))
    
    add_heading_styled(doc, "7.4 Document References", 2)
    doc_refs = app.get("references", [
        ["Enterprise Architecture Blueprint", "Confluence / Architecture / DP-005"],
        ["Security & Compliance Standard", "SecOps Policy Portal / SEC-800"]
    ])
    add_styled_table(doc, ["Title", "Location"], doc_refs, [3.2, 3.3])
    
    doc.save(output_path)
    return output_path
