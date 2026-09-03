import os
import io
import csv
from typing import Dict, Any, List

def parse_docx(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_cells:
                    full_text.append(" | ".join(row_cells))
        return "\n".join(full_text)
    except Exception as e:
        return f"[Error parsing DOCX: {str(e)}]"

def parse_pdf(file_path: str) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{text.strip()}")
        return "\n\n".join(text_parts)
    except Exception:
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_parts = []
                for i, page in enumerate(reader.pages):
                    t = page.extract_text()
                    if t:
                        text_parts.append(f"--- Page {i + 1} ---\n{t.strip()}")
                return "\n\n".join(text_parts)
        except Exception as e:
            return f"[Error parsing PDF: {str(e)}]"

def parse_excel(file_path: str) -> str:
    try:
        import pandas as pd
        excel_file = pd.ExcelFile(file_path)
        content_parts = []
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            content_parts.append(f"Sheet: {sheet_name}\n" + df.to_string(index=False))
        return "\n\n".join(content_parts)
    except Exception as e:
        return f"[Error parsing Excel: {str(e)}]"

def parse_csv(file_path: str) -> str:
    try:
        import pandas as pd
        df = pd.read_csv(file_path)
        return df.to_string(index=False)
    except Exception:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            return f"[Error parsing CSV: {str(e)}]"

def parse_image(file_path: str) -> str:
    # Try pytesseract if tesseract is installed, otherwise provide descriptive placeholder
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        if text.strip():
            return f"[OCR Extracted Text from {os.path.basename(file_path)}]:\n{text.strip()}"
    except Exception:
        pass
    
    # Fallback image description
    try:
        from PIL import Image
        img = Image.open(file_path)
        return f"[Image Process Screenshot: {os.path.basename(file_path)}, Dimensions: {img.size[0]}x{img.size[1]}px, Format: {img.format}]"
    except Exception as e:
        return f"[Screenshot / Image: {os.path.basename(file_path)}]"

def parse_text(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        return f"[Error reading file: {str(e)}]"

def extract_file_content(file_path: str, filename: str) -> Dict[str, Any]:
    ext = os.path.splitext(filename)[1].lower()
    content = ""
    
    if ext in ['.docx']:
        content = parse_docx(file_path)
    elif ext in ['.pdf']:
        content = parse_pdf(file_path)
    elif ext in ['.xlsx', '.xls']:
        content = parse_excel(file_path)
    elif ext in ['.csv']:
        content = parse_csv(file_path)
    elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.webp']:
        content = parse_image(file_path)
    elif ext in ['.txt', '.md', '.json', '.log', '.rst']:
        content = parse_text(file_path)
    else:
        content = parse_text(file_path)
        
    return {
        "filename": filename,
        "extension": ext,
        "size_bytes": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
        "content": content,
        "char_count": len(content)
    }
